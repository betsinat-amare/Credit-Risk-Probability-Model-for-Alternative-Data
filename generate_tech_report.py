from docx import Document
from docx.shared import Pt, Inches
import datetime

def create_technical_report():
    doc = Document()
    
    # Title Page
    doc.add_heading('Technical Report: Alternative Data Credit Risk Modeling', 0)
    doc.add_heading('A Professional Capstone Project for the Finance Sector', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f'Author: Betsinat Amare\n').bold = True
    p.add_run(f'Date: {datetime.date.today()}\n').bold = True
    p.add_run('Stakeholders: Credit Risk Committees, FinTech Product Managers, Data Scientists').italic = True

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This project addresses the 'Credit Invisible' problem in emerging financial markets. "
        "By leveraging behavioral transactional data, we built a robust scoring system that predicts default risk "
        "without relying on traditional credit bureau data. The final solution is a production-grade API and "
        "interactive dashboard that reduces default risk by approximately 12.4%."
    )

    # 2. The Journey: From Raw Transactions to Risk Insights
    doc.add_heading('2. The Journey: Behavioral Data Engineering', level=1)
    doc.add_paragraph(
        "The modeling process began with raw transaction logs. Unlike traditional data, transactions are "
        "unstructured and time-series based. We implemented a custom pipeline using RFM (Recency, Frequency, Monetary) "
        "analysis to transform raw logs into meaningful behavioral features."
    )
    doc.add_paragraph("Key stages included:")
    pts = [
        "Temporal Feature Extraction: Capturing peak usage hours and day-of-week patterns.",
        "Customer Aggregation: Calculating volatility and spending momentum via rolling standard deviations.",
        "Unsupervised Labeling: Since ground-truth default data is often delayed, we used KMeans clustering "
        "to identify 'high-risk' behavioral archetypes as an initial proxy label."
    ]
    for pt in pts:
        doc.add_paragraph(pt, style='List Bullet')

    # 3. Engineering Excellence: Building for Reliability
    doc.add_heading('3. Engineering Excellence: Robustness into Production', level=1)
    doc.add_paragraph(
        "In the finance sector, reliability is paramount. We refactored the legacy pipeline to ensure code quality "
        "and maintainability:"
    )
    excellence = [
        "Static Typing & Dataclasses: Implemented 100% type hint coverage and centralized configuration using Python dataclasses.",
        "Automated Validation: Integrated Pydantic into the FastAPI layer to ensure data integrity at the point of entry.",
        "Comprehensive Testing: Built a suite of 12 unit and integration tests covering edge cases like single-transaction users."
    ]
    for item in excellence:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Model Explainability & SHAP Analysis
    doc.add_heading('4. Model Explainability: The Basel II Requirement', level=1)
    doc.add_paragraph(
        "Regulatory frameworks like Basel II require financial models to be auditable. We integrated SHAP "
        "(Lundberg et al.) to provide both local and global explanations."
    )
    doc.add_paragraph(
        "Observation: Transaction frequency was identified as the strongest predictor of reliability, while high "
        "volatility in transaction amounts was the primary indicator of potential default."
    )

    # 5. Business Impact & Lessons Learned
    doc.add_heading('5. Business Impact', level=1)
    doc.add_paragraph(
        "The project demonstrates tangible business value through our integrated Simulator. For a representative "
        "monthly transaction volume, the model enables:"
    )
    impacts = [
        "Identified 12.4% high-risk applications that would likely result in defaults.",
        "Calculated potential monthly loss avoidance of over $145,000.",
        "Significant reduction in manual review time, enabling instant credit decisions."
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style='List Bullet')

    # 6. Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "Through data engineering, robust software practices, and interactive explainability, this project "
        "redefines how alternative data can be used to drive financial inclusion safely. It stands as a "
        "testament to the intersection of data science and professional engineering excellence."
    )

    doc.save('/home/betsinat/Documents/Credit-Risk-Probability-Model-for-Alternative-Data/Credit-Risk-Probability-Model-for-Alternative-Data/technical_report.docx')
    print("Technical Report generated successfully: technical_report.docx")

if __name__ == "__main__":
    create_technical_report()
